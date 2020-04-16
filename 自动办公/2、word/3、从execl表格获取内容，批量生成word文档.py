import xlrd
from docx import Document

def chang_text(old_text,new_text):
    all_paragraphs = document.paragraphs
    for all_paragraph in all_paragraphs:
        for run in paragraph.runs:
            run_text = run.text.replace(old_text,new_text)
            run.text = run_text

    all_tables = document.tables
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.replace(old_text,new_text)
                cell.text = cell_text

xlsx = xlrd.open_workbook('合同信息表.xlsx')
sheet = xlsx.sheet_by_index(0)

for table_row in range(1,sheet.nrows):
    document = Document('模板.docx')
    for table_col in range(0,sheet.ncols):
        chang_text(str(sheet.cell_value(0,table_col)),str(sheet.cell_value(table_row,table_col)))

    document.save("%s合同.docx"% str(sheet.cell_value(table_row,0)))
    print("%s合同完成"% str(sheet.cell_value(table_row,0)))




