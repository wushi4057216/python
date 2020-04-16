from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
import time

company_list = ['客户1', '客户2', '客户3', ]
today = time.strftime("%Y{y}%m{m}%d{d}", time.localtime()).format(y='年', m='月', d='日', )
#获取今日时间，整理成“年-月-日”格式

for i in company_list:
    document = Document()
    document.styles["Normal"].font.name = u'宋体' #设置文档默认字体为宋体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') #设置中文字体
    p1 = document.add_paragraph() #初始化自然段
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER   #对齐设置成居中，默认为左对齐

    run1 = p1.add_run('关于下达%s的通知'%(today))
    run1.font.name = '微软雅黑'  #设置段落字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run1.font.size = Pt(21)  #设置字体大小
    run1.font.bold = True #设置加粗
    p1.space_after = Pt(5) #段后距离5磅
    p1.space_before = Pt(5) #段前距离5磅

    p2 = document.add_paragraph()
    run2 = p2.add_run(i + ':')
  #  run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run2.font.size = Pt(16)  #设置字体大小
    run2.font.bold = True #设置加粗


    p3 = document.add_paragraph()
    run3 = p3.add_run('  根据公司安排，XXXXX')
  #  run3._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run3.font.size = Pt(16)  #设置字体大小
    run3.font.bold = True #设置加粗


    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('联系人：小杨 联系电话：1234567890')
   # run4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run4.font.size = Pt(16)  #设置字体大小
    run4.font.bold = True #设置加粗



    document.save('%s产品文档.doc'%i)















