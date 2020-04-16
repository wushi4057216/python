from docx import Document
import xlrd


def change_text(old_text, new_text):
    all_paragraphs = document.paragraphs
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            run_text = run.text.replace(old_text, new_text)
            run.text = run_text

    all_tables = document.tables
    for table in all_tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.replace(old_text, new_text)
                cell.text = cell_text


xlsx = xlrd.open_workbook('d:/合同信息表.xlsx')
sheet = xlsx.sheet_by_index(0)

for table_row in range(1, sheet.nrows):
    document = Document("d:/修改模板.docx")
    for table_col in range(0, sheet.ncols):
        change_text(str(sheet.cell_value(0, table_col)), str(sheet.cell_value(table_row, table_col)))

    document.save("d:/%s合同.docx" % str(sheet.cell_value(table_row, 2)))
    print("%s合同完成" % str(sheet.cell_value(table_row, 2)))
