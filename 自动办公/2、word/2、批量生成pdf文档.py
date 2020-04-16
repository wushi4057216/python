from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
from docx.shared import Inches #修改图片尺寸
#以上是docx库
from win32com.client import Dispatch, constants, gencache
#安装使用pip install pywin32
import datetime
import time
import os

today = time.strftime("%y{y}%m{m}%d{d}", time.localtime()).format(y='年', m='月', d='日')
company_list = ['客户1', '客户2', '客户3', ]

for i in company_list:
    document = Document()
    document.styles["Normal"].font.name = u'宋体' #设置文档默认字体为宋体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') #设置中文字体

    document.add_picture('b.jpg', width=Inches(6))#添加图片，Inches设置图片大小

    #设置第一自然段内容
    p1 = document.add_paragraph() #初始化自然段
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER   #对齐设置成居中，默认为左对齐
    run1 = p1.add_run('关于下达%s的通知'%(today))
    run1.font.name = '微软雅黑'  #设置段落字体
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run1.font.size = Pt(21)  #设置字体大小
    run1.font.bold = True #设置加粗
    p1.space_after = Pt(5) #段后距离5磅
    p1.space_before = Pt(5) #段前距离5磅

    #添加表格
    table = document.add_table(rows=3,cols=3,style='Table Grid')#初始化一个表格，3行3列，表格类型为Table Grid
    #合并单元格从0行0列到0行2列合并
    table.cell(0, 0).merge(table.cell(0, 2))
    #设置单元格字体，paragraphs是段落
    table_run1 = table.cell(0, 0).paragraphs[0].add_run('产品价格表')
    table_run1.font.name = '微软雅黑'  #设置段落字体
    table_run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.cell(1, 0).text = "日期"  #填入单元格内容


    #换页，插入分页符
    document.add_page_break()

    p2 = document.add_paragraph()  # 初始化自然段
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 对齐设置成居中，默认为左对齐
    run2 = p2.add_run('此处是广告')

    #保存文档,判断是否有重名文件
    if os.path.exists('%s-价格表.docx' %i):
        os.remove('%s-价格表.docx' %i)
    document.save('%s-价格表.docx' %i)
    doc_path = '%s-价格表.docx' %i
    pdf_path = '%s-价格表.pdf' %i

    gencache.EnsureModule('{00020905-0000-0000-C000-00000000046}', 0, 8, 4)

    wd = Dispatch('Word.Application')

    doc = wd.Documents.Open(doc_path, ReadOnly=1)
    doc.ExportAsFixedFormat(pdf_path, constants.wdEnportFormatPDF, Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    wd.Quit(constants.wdDoNotSaveChanges)

    time.sleep(5)
    os.remove('%s-价格表.docx' %i)














