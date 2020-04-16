import time
import datetime
import os
from docx import Document  # 导入库
from docx.shared import Inches
from selenium import webdriver

#定义时间
today = time.strftime("%Y-%m-%d",time.localtime())

###定义时间值到变量，
x = datetime.datetime.now() #現在時間
year=str(x.year)
month=str(x.month)
day=str(x.day)
hour = str(x.hour)
minute = int(x.minute)
if minute >=30:
    minute = '30'
    #print(minute)
else:
    minute ='00'
    #print(minute)
hourm = str(hour + '-' + minute)
print(hourm)
fileyear='d:/'+year
filemonth=fileyear+'/'+month
fileday=filemonth+'/'+day
filehour=fileday+'/'+hourm
#print(filehourm)
#判断创建现在日期的逐层目录。
if not os.path.exists(fileyear):
    os.mkdir(fileyear)
    os.mkdir(filemonth)
    os.mkdir(fileday)
    os.mkdir(filehour)
else:
    if not os.path.exists(filemonth):
        os.mkdir(filemonth)
        os.mkdir(fileday)
        os.mkdir(filehour)
    else:
        if not os.path.exists(fileday):
            os.mkdir(fileday)
            os.mkdir(filehour)
        else:
            if not os.path.exists(filehour):
                os.mkdir(filehour)

#读取表格内容，获取名称和网址
path = "“护网2019”期间监测报告每日汇报表模板.docx" #文件路径
document = Document(path) #读入文件
tables = document.tables #获取文件中的表格集
table = tables[2]#获取文件中的第三个表格

#定义使用浏览器
driver = webdriver.Chrome()
#driver = webdriver.Chrome(executable_path="C:\Program Files (x86)\Google\Chrome\Application\chromedriver.exe")

def wz(table):
###从时间表格第二行开始循环读取表格数据到变量
    for i in range(1,len(table.rows)):
        webname = table.cell(i,0).text
        weburl = table.cell(i,2).text
        jietu = table.cell(i, 6).text
        print(webname)
        print(weburl)
        print(jietu)
    ###排除空行表格，空白不输出
       # if  name1.split():
       #     name = name1
       # if  web1.split():
       #     web = web1

    ###浏览器操作
        # driver.refresh()     # 刷新页面
        # # driver.set_window_size(560, 960, CURRENT)     # 设置屏幕尺寸
        driver.get(weburl)  # 打开web页面
        time.sleep(3)
        # #driver.maximize_window() # 最大化窗口
        driver.get_screenshot_as_file(filehour + '/' + webname + '.png') # 保存浏览器截图并保存到fileday

    ###输出内容到表格第5列，从第二行开始
        cell = table.cell(i, 6)
        ph = cell.paragraphs[0]
        run = ph.add_run('')
      #  run.add_break()  # 添加一个折行
        run.add_picture(filehour + '/' + webname + '.png',width=Inches(3.0))  # 插入图像，可以是内存中的图像，width=Inches(1.0)指定宽度。



###判断当前时间点，写入到对应时间表格
if hourm == "8-00":
    table = tables[2]
    wz(table)
if hourm == "8-30":
    table = tables[3]
    wz(table)
if hourm == "9-00":
    table = tables[4]
    wz(table)
if hourm == "9-30":
    table = tables[5]
    wz(table)
if hourm == "10-00":
    table = tables[6]
    wz(table)
if hourm == "10-30":
    table = tables[7]
    wz(table)
if hourm == "11-00":
    table = tables[8]
    wz(table)
if hourm == "14-00":
    table = tables[14]
    wz(table)
if hourm == "14-30":
    table = tables[15]
    wz(table)
if hourm == "15-00":
    table = tables[16]
    wz(table)
if hourm == "15-30":
    table = tables[17]
    wz(table)
if hourm == "16-00":
    table = tables[18]
    wz(table)
if hourm == "16-30":
    table = tables[19]
    wz(table)
if hourm == "17-00":
    table = tables[20]
    wz(table)
if hourm == "17-30":
    table = tables[21]
    wz(table)
if hourm == "21-30":
    table = tables[21]
    wz(table)
if hourm == "11-30":
    for zw in [9,10,11,12,13]:
        table = tables[zw]
        wz(table)


document.save("“护网2019”期间监测报告每日汇报表模板.docx")    #另存文档
driver.quit() #退出浏览器，close()是关闭当前访问页面，quit()是退出浏览器，结束进程，且回收临时文件
exit()