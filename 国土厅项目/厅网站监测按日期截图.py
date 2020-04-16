import time
import datetime
import os
from docx import Document  # 导入库
from docx.shared import Inches
from selenium import webdriver

#定义时间
today = time.strftime("%Y-%m-%d",time.localtime())

###定义时间值到变量，
x = datetime.datetime.now() #現在時間
year=str(x.year)
month=str(x.month)
day=str(x.day)
hour = str(x.hour)
minute = int(x.minute)
if minute >=30:
    minute = '30'
    print(minute)
else:
    minute ='00'
    print(minute)
hourm = str(hour + '-' + minute)
print(hourm)
fileyear='d:/'+year
filemonth=fileyear+'/'+month
fileday=filemonth+'/'+day
filehour=fileday+'/'+hourm
print(filehour)
#判断创建现在日期的逐层目录。
if not os.path.exists(fileyear):
    os.mkdir(fileyear)
    os.mkdir(filemonth)
    os.mkdir(fileday)
    os.mkdir(filehour)
else:
    if not os.path.exists(filemonth):
        os.mkdir(filemonth)
        os.mkdir(fileday)
        os.mkdir(filehour)
    else:
        if not os.path.exists(fileday):
            os.mkdir(fileday)
            os.mkdir(filehour)
        else:
            if not os.path.exists(filehour):
                os.mkdir(filehour)

#读取表格内容，获取名称和网址
path = "业务可用性检查报告模板.docx" #文件路径
document = Document(path) #读入文件
tables = document.tables #获取文件中的表格集
table1 = tables[0]#获取文件中的第一个表格

#定义使用浏览器
driver = webdriver.Chrome()
#driver = webdriver.Chrome(executable_path="C:\Program Files (x86)\Google\Chrome\Application\chromedriver.exe")

for i in range(1,20):#从表格第二行开始循环读取表格数据到天地图
    webname = table1.cell(i,1).text
    weburl = table1.cell(i,2).text

###排除空行表格，空白不输出
   # if  name1.split():
   #     name = name1
   # if  web1.split():
   #     web = web1

###浏览器操作
    # driver.refresh()     # 刷新页面
    # # driver.set_window_size(560, 960, CURRENT)     # 设置屏幕尺寸
    driver.get(weburl)  # 打开web页面
    time.sleep(3)
    # #driver.maximize_window() # 最大化窗口
    driver.get_screenshot_as_file(filehour + '/' + webname + '.png') # 保存浏览器截图并保存到fileday

###输出内容到表格第二列，从第二行开始
    cell = table2.cell(i, 1)
    ph = cell.paragraphs[0]
    run = ph.add_run('')
  #  run.add_break()  # 添加一个折行
    run.add_picture(filehour + '/' + webname + '.png', width=Inches(3.0))  # 插入图像，可以是内存中的图像，width=Inches(1.0)指定宽度。

document.save('业务可用性检查报告'+today+'.docx')    #另存文档
driver.close() #退出浏览器，close()是关闭当前访问页面，quit()是退出浏览器，结束进程，且回收临时文件

