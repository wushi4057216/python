from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

document = Document()
document.add_heading('This is my title', 0)
document.add_paragraph('my paragraph')

document.styles['Normal'].font.name = u'黑体'
p = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u'我添加的段落文字 ')
run.font.color.rgb = RGBColor(54, 95, 145)
run.font.size = Pt(36)

#pic = document.add_picture('logo1.PNG')
#last_paragraph = document.paragraphs[-1]
#last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

rows = 2
cols = 3
table = document.add_table(rows=rows, cols=cols, style="Table Grid")  # 添加2行3列的表格

for i in range(rows):
    tr = table.rows[i]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "450")
    trPr.append(trHeight)  # 表格高度设置
# table.autofit = False
col = table.columns[1]
col.width = Inches(5)
arr = [u'序号', u"类型", u"详细描述"]
heading_cells = table.rows[0].cells
for i in range(cols):
    p = heading_cells[i].paragraphs[0]
    run = p.add_run(arr[i])
    run.font.color.rgb = RGBColor(54, 95, 145)  # 颜色设置，这里是用RGB颜色
    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.cell(1, 1).text = u'表格文字'
table.add_row()
document.save('test1.docx')
